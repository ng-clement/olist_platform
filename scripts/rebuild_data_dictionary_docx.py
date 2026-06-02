"""Rebuild OLP-DD-001_Data_Dictionary.docx with full column tables."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0D, 0x3B, 0x66)
TEAL = RGBColor(0x1C, 0x72, 0x93)
LIGHT_BLUE = RGBColor(0xD6, 0xEA, 0xF8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0xF4, 0xF6, 0xF7)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OLIST MODERN DATA PLATFORM")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Data Dictionary")
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = TEAL

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Complete Column Reference for All Warehouse Tables and Marts")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = [
        ("Document ID", "OLP-DD-001"),
        ("Version", "1.0"),
        ("Status", "Final"),
        ("Date", "May 2026"),
        ("Classification", "Internal — Confidential"),
        ("Owner", "Data Engineering Team"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[0].width = Inches(2)
        t.rows[i].cells[1].width = Inches(4)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else TEAL


def add_column_table(doc, headers, rows):
    """Add a professional column-definition table."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        shade_cell(cell, "0D3B66")
        p = cell.paragraphs[0]
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = WHITE
        p.runs[0].font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.add_row()
        fill = "EBF5FB" if ri % 2 == 0 else "FDFEFE"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            shade_cell(cell, fill)
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)
    doc.add_paragraph()


def build():
    doc = Document()
    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_title_block(doc)

    # ── Section 1: Overview ────────────────────────────────────────────────────
    add_heading(doc, "1.  Overview")
    doc.add_paragraph(
        "This Data Dictionary documents every column in the Olist Analytics Platform BigQuery datasets: "
        "olist_analytics (star schema dimensions and facts), olist_analytics_staging (dbt staging views), "
        "and olist_analytics_marts (dbt mart tables).\n\n"
        "Naming conventions: dimension tables prefixed with Dim; fact tables prefixed with Fact; "
        "staging views prefixed with stg_; mart tables prefixed with mart_.\n\n"
        "All tables include a dw_inserted_at TIMESTAMP column recording the pipeline run timestamp. "
        "Surrogate keys use FARM_FINGERPRINT() on natural keys for deterministic generation."
    )

    # ── Section 2: Source Datasets ─────────────────────────────────────────────
    add_heading(doc, "2.  Source Datasets (olist_raw)")
    doc.add_paragraph(
        "Raw tables in the olist_raw dataset are written by Meltano (tap-csv → target-bigquery) "
        "with WRITE_TRUNCATE on each pipeline run. These tables are the inputs to all dbt staging models."
    )
    add_column_table(
        doc,
        ["Table", "Source CSV", "Rows", "Description"],
        [
            (
                "orders",
                "olist_orders_dataset.csv",
                "99,441",
                "All orders across all sellers, 2016–2018",
            ),
            (
                "order_items",
                "olist_order_items_dataset.csv",
                "112,650",
                "Line-item detail per order",
            ),
            (
                "customers",
                "olist_customers_dataset.csv",
                "99,441",
                "One row per order–customer pair (not per unique person)",
            ),
            (
                "products",
                "olist_products_dataset.csv",
                "32,951",
                "Product catalogue with dimensions",
            ),
            (
                "sellers",
                "olist_sellers_dataset.csv",
                "3,095",
                "Seller accounts on the platform",
            ),
            (
                "order_payments",
                "olist_order_payments_dataset.csv",
                "103,886",
                "Payment records (multiple per order)",
            ),
            (
                "order_reviews",
                "olist_order_reviews_dataset.csv",
                "99,224",
                "Customer satisfaction reviews",
            ),
            (
                "geolocation",
                "olist_geolocation_dataset.csv",
                "1,000,163",
                "GPS readings per zip code prefix",
            ),
            (
                "marketing_qualified_leads",
                "olist_marketing_qualified_leads_dataset.csv",
                "8,000",
                "B2B seller acquisition funnel",
            ),
            (
                "closed_deals",
                "olist_closed_deals_dataset.csv",
                "842",
                "Converted sellers (10.5% of MQLs)",
            ),
            (
                "product_category_translation",
                "product_category_name_translation.csv",
                "71",
                "Portuguese → English category names",
            ),
        ],
    )

    # ── Section 3: Staging Views ───────────────────────────────────────────────
    add_heading(doc, "3.  Staging Views (olist_analytics_staging)")
    doc.add_paragraph(
        "Staging models are BigQuery views — zero storage cost; they always reflect the latest raw data. "
        "Each view applies type-casting, null-filtering on primary keys, and string normalisation."
    )

    add_heading(doc, "3.1  stg_orders", level=2)
    doc.add_paragraph("Source: olist_raw.orders  |  Grain: one row per order_id")
    add_column_table(
        doc,
        ["Column", "Type", "Transformation / Notes"],
        [
            ("order_id", "STRING", "Primary key; validated unique + not_null"),
            ("customer_id", "STRING", "FK → stg_customers"),
            (
                "order_status",
                "STRING",
                "Accepted: delivered, shipped, canceled, unavailable, invoiced, processing, created, approved",
            ),
            (
                "purchase_date",
                "DATE",
                "SAFE_CAST from order_purchase_timestamp TIMESTAMP",
            ),
            ("approved_date", "DATE", "SAFE_CAST from order_approved_at"),
            ("carrier_date", "DATE", "SAFE_CAST from order_delivered_carrier_date"),
            ("delivered_date", "DATE", "SAFE_CAST from order_delivered_customer_date"),
            (
                "estimated_delivery_date",
                "DATE",
                "SAFE_CAST from order_estimated_delivery_date",
            ),
            (
                "delivery_days",
                "INT64",
                "DATE_DIFF(delivered_date, purchase_date, DAY); NULL if undelivered",
            ),
            (
                "is_on_time",
                "BOOL",
                "delivered_date ≤ estimated_delivery_date; NULL if undelivered",
            ),
        ],
    )

    add_heading(doc, "3.2  stg_order_items", level=2)
    doc.add_paragraph(
        "Source: olist_raw.order_items  |  Grain: one row per (order_id, order_item_id)"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Transformation / Notes"],
        [
            ("order_id", "STRING", "FK → stg_orders; validated not_null"),
            ("order_item_id", "INT64", "Sequence number within the order"),
            ("product_id", "STRING", "FK → stg_products; validated not_null"),
            ("seller_id", "STRING", "FK → stg_sellers; validated not_null"),
            ("shipping_limit_date", "DATE", "SAFE_CAST from timestamp"),
            ("item_price", "FLOAT64", "Validated > 0"),
            ("freight_value", "FLOAT64", "Validated ≥ 0"),
            ("item_total", "FLOAT64", "item_price + freight_value"),
        ],
    )

    add_heading(doc, "3.3  stg_customers", level=2)
    doc.add_paragraph(
        "Source: olist_raw.customers  |  Grain: one row per customer_id (order-scoped)"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Transformation / Notes"],
        [
            (
                "customer_id",
                "STRING",
                "Primary key (unique per order); validated unique + not_null",
            ),
            (
                "customer_unique_id",
                "STRING",
                "Stable person-level identifier; validated not_null",
            ),
            ("zip_code_prefix", "INT64", "CAST from string"),
            ("customer_city", "STRING", "LOWER(TRIM(…))"),
            (
                "customer_state",
                "STRING",
                "UPPER(TRIM(…)); accepted: all 27 Brazilian state codes",
            ),
        ],
    )

    add_heading(doc, "3.4  stg_products", level=2)
    doc.add_paragraph(
        "Source: olist_raw.products LEFT JOIN olist_raw.product_category_translation  |  Grain: one row per product_id"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Transformation / Notes"],
        [
            ("product_id", "STRING", "Primary key; validated unique + not_null"),
            ("product_category_name", "STRING", "Original Portuguese category name"),
            (
                "category_english",
                "STRING",
                "COALESCE(English translation, Portuguese name, 'Unknown') — triple fallback ensures no NULLs",
            ),
            (
                "product_name_length",
                "INT64",
                "Character count (renamed from product_name_lenght)",
            ),
            (
                "product_description_length",
                "INT64",
                "Character count (renamed from product_description_lenght)",
            ),
            ("product_photos_qty", "INT64", "Number of product photos"),
            ("product_weight_g", "INT64", "Weight in grams"),
            (
                "product_length_cm / _height_cm / _width_cm",
                "FLOAT64",
                "Physical dimensions",
            ),
            ("volume_litres", "FLOAT64", "length × height × width / 1,000,000"),
        ],
    )

    add_heading(doc, "3.5  stg_sellers", level=2)
    doc.add_paragraph("Source: olist_raw.sellers  |  Grain: one row per seller_id")
    add_column_table(
        doc,
        ["Column", "Type", "Transformation / Notes"],
        [
            ("seller_id", "STRING", "Primary key; validated unique + not_null"),
            ("zip_code_prefix", "INT64", "CAST from string"),
            ("seller_city", "STRING", "LOWER(TRIM(…))"),
            (
                "seller_state",
                "STRING",
                "UPPER(TRIM(…)); validated accepted state codes",
            ),
        ],
    )

    add_heading(doc, "3.6  stg_payments", level=2)
    doc.add_paragraph(
        "Source: olist_raw.order_payments  |  Grain: one row per (order_id, payment_sequential)  |  Filter: excludes payment_type = 'not_defined'"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Transformation / Notes"],
        [
            ("order_id", "STRING", "FK → stg_orders; validated not_null"),
            (
                "payment_sequential",
                "INT64",
                "Sequence number (1 = primary payment method)",
            ),
            (
                "payment_type",
                "STRING",
                "Accepted: credit_card, boleto, voucher, debit_card",
            ),
            ("payment_installments", "INT64", "Number of instalment periods"),
            ("payment_value", "FLOAT64", "Validated ≥ 0"),
        ],
    )

    add_heading(doc, "3.7  stg_reviews", level=2)
    doc.add_paragraph(
        "Source: olist_raw.order_reviews  |  Grain: one row per review_id"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Transformation / Notes"],
        [
            ("review_id", "STRING", "Review identifier; validated not_null"),
            ("order_id", "STRING", "FK → stg_orders; validated not_null"),
            ("review_score", "INT64", "Validated 1–5 via accepted_range"),
            ("review_comment_title", "STRING", "Optional; may be NULL"),
            ("review_comment_message", "STRING", "Optional; may be NULL"),
            ("review_creation_date", "DATE", "SAFE_CAST from timestamp"),
            ("review_answer_timestamp", "TIMESTAMP", "SAFE_CAST"),
        ],
    )

    add_heading(doc, "3.8  stg_marketing_leads", level=2)
    doc.add_paragraph(
        "Source: olist_raw.marketing_qualified_leads LEFT JOIN olist_raw.closed_deals  |  Grain: one row per mql_id"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Transformation / Notes"],
        [
            ("mql_id", "STRING", "Primary key; validated unique + not_null"),
            ("landing_page_id", "STRING", "Landing page of origin"),
            (
                "origin",
                "STRING",
                "Acquisition channel; ~2% NULL (severity: warn — direct/organic traffic)",
            ),
            ("first_contact_date", "DATE", "SAFE_CAST"),
            ("won_date", "DATE", "From closed_deals; NULL if not converted"),
            ("seller_id", "STRING", "From closed_deals; NULL if not converted"),
            ("business_segment", "STRING", "From closed_deals"),
            ("lead_type", "STRING", "From closed_deals"),
            ("lead_behaviour_profile", "STRING", "From closed_deals"),
            ("business_type", "STRING", "reseller, manufacturer, other"),
            (
                "declared_monthly_revenue",
                "FLOAT64",
                "Self-reported; NULL if not provided",
            ),
            ("is_converted", "BOOL", "TRUE when a closed deal with a seller_id exists"),
            (
                "days_to_close",
                "INT64",
                "DATE_DIFF(won_date, first_contact_date, DAY); NULL if not converted",
            ),
        ],
    )

    add_heading(doc, "3.9  stg_geolocation", level=2)
    doc.add_paragraph(
        "Source: olist_raw.geolocation (1,000,163 rows)  |  Grain: one row per zip_code_prefix (19,015 rows after deduplication via median lat/lng, mode city/state)"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Transformation / Notes"],
        [
            (
                "zip_code_prefix",
                "INT64",
                "5-digit CEP prefix; validated unique, 1001–99999",
            ),
            (
                "latitude",
                "FLOAT64",
                "Median across all GPS readings; validated −35.0 to 5.3",
            ),
            (
                "longitude",
                "FLOAT64",
                "Median across all GPS readings; validated −74.0 to −28.0",
            ),
            ("city_normalized", "STRING", "Mode city name (lower-case)"),
            (
                "city",
                "STRING",
                "INITCAP(city_normalized) with hyphen→space normalisation",
            ),
            (
                "state_code",
                "STRING",
                "Mode state; validated against 27 Brazilian state codes",
            ),
            (
                "region",
                "STRING",
                "Derived from state_code: Southeast, South, Midwest, Northeast, North",
            ),
            (
                "raw_row_count",
                "INT64",
                "Number of raw GPS readings aggregated into this prefix",
            ),
        ],
    )

    # ── Section 4: Dimension Tables ────────────────────────────────────────────
    add_heading(doc, "4.  Dimension Tables (olist_analytics)")

    add_heading(doc, "4.1  DimDate", level=2)
    doc.add_paragraph(
        "Grain: one row per calendar day, 2016-01-01 to 2020-12-31 (1,827 rows)  |  Purpose: time-based slicing across all fact tables"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("full_date", "DATE", "Calendar date (primary key)"),
            (
                "date_key",
                "STRING",
                "Date formatted as YYYYMMDD (surrogate for legacy BI tools)",
            ),
            ("year", "INT64", "Calendar year"),
            ("quarter", "INT64", "Quarter (1–4)"),
            ("month", "INT64", "Month number (1–12)"),
            ("month_name", "STRING", "Full month name (e.g., 'January')"),
            ("month_abbr", "STRING", "3-letter abbreviation"),
            ("week_of_year", "INT64", "ISO week number"),
            ("day_of_month", "INT64", "Day within the month (1–31)"),
            ("day_of_week", "INT64", "Day of week (1=Sunday, 7=Saturday)"),
            ("day_name", "STRING", "Full day name (e.g., 'Monday')"),
            ("is_weekday", "BOOL", "TRUE if Mon–Fri"),
            ("is_public_holiday", "BOOL", "TRUE for Brazilian fixed public holidays"),
            ("first_day_of_month", "DATE", "First day of the calendar month"),
            ("last_day_of_month", "DATE", "Last day of the calendar month"),
            ("first_day_of_quarter", "DATE", "First day of the calendar quarter"),
            ("half_year_label", "STRING", "Half-year label (e.g., 'H1-2017')"),
        ],
    )

    add_heading(doc, "4.2  DimGeography", level=2)
    doc.add_paragraph(
        "Grain: one row per zip_code_prefix (19,015 rows)  |  Source: deduplicated from 1,000,163 raw GPS readings"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("geo_key", "INT64", "Surrogate key — FARM_FINGERPRINT(zip_code_prefix)"),
            ("zip_code_prefix", "INT64", "5-digit Brazilian CEP prefix (natural key)"),
            ("zip_code_formatted", "STRING", "Zero-padded string (e.g., '01001')"),
            (
                "latitude",
                "FLOAT64",
                "Median latitude across all GPS readings for this prefix",
            ),
            ("longitude", "FLOAT64", "Median longitude across all GPS readings"),
            ("geo_point", "GEOGRAPHY", "BigQuery GEOGRAPHY point for spatial queries"),
            ("city", "STRING", "Most common city name (INITCAP normalised)"),
            ("city_normalized", "STRING", "Lower-case city name"),
            ("state_code", "STRING", "2-letter Brazilian state code (e.g., 'SP')"),
            ("state_name", "STRING", "Full state name (e.g., 'São Paulo')"),
            (
                "region",
                "STRING",
                "Brazilian macro-region: Southeast, South, Midwest, Northeast, North",
            ),
            (
                "is_frontier_market",
                "BOOL",
                "FALSE for 10 low-penetration states: AC, AL, AM, AP, MA, PB, RN, RO, RR, TO",
            ),
            (
                "geographic_zone",
                "STRING",
                "Coastal, Amazon Basin, Central Plateau, Southern Cone",
            ),
            (
                "raw_row_count",
                "INT64",
                "Number of raw GPS readings aggregated into this prefix",
            ),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    add_heading(doc, "4.3  DimCustomer", level=2)
    doc.add_paragraph(
        "Grain: one row per customer_id (99,441 rows)  |  Note: customer_id is order-specific. Use customer_unique_id for person-level analysis."
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            (
                "customer_id",
                "STRING",
                "Order-specific customer identifier (natural key)",
            ),
            (
                "customer_unique_id",
                "STRING",
                "Stable person-level identifier; one person = one unique_id across multiple orders",
            ),
            ("zip_code_prefix", "INT64", "Customer's 5-digit zip code prefix"),
            ("city", "STRING", "Customer city (lower-case normalised)"),
            ("state_code", "STRING", "2-letter state code"),
            ("state_name", "STRING", "Full state name (from DimGeography join)"),
            ("region", "STRING", "Brazilian macro-region"),
            ("customer_lat", "FLOAT64", "Approximate latitude from DimGeography"),
            ("customer_lng", "FLOAT64", "Approximate longitude from DimGeography"),
            ("geographic_zone", "STRING", "Geographic zone from DimGeography"),
            (
                "is_frontier_market",
                "BOOL",
                "Whether customer is in a high-penetration market",
            ),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    add_heading(doc, "4.4  DimProduct", level=2)
    doc.add_paragraph("Grain: one row per product_id (32,951 rows)")
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("product_id", "STRING", "Unique product identifier (natural key)"),
            ("product_category_name", "STRING", "Portuguese category name"),
            (
                "category_english",
                "STRING",
                "English category name; falls back to Portuguese name, then 'Unknown'",
            ),
            ("product_name_length", "INT64", "Character count of product name"),
            (
                "product_description_length",
                "INT64",
                "Character count of product description",
            ),
            ("product_photos_qty", "INT64", "Number of product photos"),
            ("product_weight_g", "INT64", "Weight in grams"),
            ("product_length_cm", "FLOAT64", "Length in cm"),
            ("product_height_cm", "FLOAT64", "Height in cm"),
            ("product_width_cm", "FLOAT64", "Width in cm"),
            (
                "volume_litres",
                "FLOAT64",
                "Volume = length × height × width / 1,000,000",
            ),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    add_heading(doc, "4.5  DimSeller", level=2)
    doc.add_paragraph("Grain: one row per seller_id (3,095 rows)")
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("seller_id", "STRING", "Unique seller identifier (natural key)"),
            ("zip_code_prefix", "INT64", "Seller's 5-digit zip code prefix"),
            ("city", "STRING", "Seller city (lower-case normalised)"),
            ("state_code", "STRING", "2-letter state code"),
            ("state_name", "STRING", "Full state name"),
            ("region", "STRING", "Brazilian macro-region"),
            ("seller_lat", "FLOAT64", "Approximate latitude from DimGeography"),
            ("seller_lng", "FLOAT64", "Approximate longitude"),
            ("geographic_zone", "STRING", "Geographic zone"),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    add_heading(doc, "4.6  DimPaymentType", level=2)
    doc.add_paragraph("Grain: one row per payment type code (4 rows)")
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("payment_type_key", "INT64", "Surrogate key"),
            (
                "payment_type_code",
                "STRING",
                "Raw code from source (e.g., 'credit_card')",
            ),
            (
                "payment_type_name",
                "STRING",
                "Human-readable name (e.g., 'Credit Card')",
            ),
            ("payment_category", "STRING", "Card, Bank Transfer, Voucher, Other"),
            ("supports_installments", "BOOL", "TRUE for credit_card and debit_card"),
            (
                "is_offline_payment",
                "BOOL",
                "TRUE for boleto (cash-in-advance bank slip)",
            ),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    add_heading(doc, "4.7  DimMarketingChannel", level=2)
    doc.add_paragraph("Grain: one row per distinct origin channel value")
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("channel_key", "INT64", "Surrogate key"),
            ("channel_code", "STRING", "Raw origin value (e.g., 'paid_search')"),
            ("channel_name", "STRING", "Display name (e.g., 'Paid Search')"),
            ("channel_type", "STRING", "Paid, Organic, Owned, Earned, Other"),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    # ── Section 5: Fact Tables ─────────────────────────────────────────────────
    add_heading(doc, "5.  Fact Tables (olist_analytics)")

    add_heading(doc, "5.1  FactOrders", level=2)
    doc.add_paragraph(
        "Grain: one row per order_id (99,441 rows)  |  Partition: purchase_date (DATE)  |  Cluster: customer_state, order_status"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("order_id", "STRING", "Unique order identifier (natural key)"),
            ("customer_id", "STRING", "FK → DimCustomer"),
            ("purchase_date", "DATE", "Date the order was placed"),
            ("approved_date", "DATE", "Date payment was approved"),
            ("carrier_date", "DATE", "Date the order was handed to the carrier"),
            ("delivered_date", "DATE", "Date the customer received the order"),
            (
                "estimated_delivery_date",
                "DATE",
                "Estimated delivery date shown at purchase",
            ),
            (
                "order_status",
                "STRING",
                "delivered, shipped, canceled, unavailable, invoiced, processing, created, approved",
            ),
            ("item_count", "INT64", "Number of items in the order"),
            ("distinct_products", "INT64", "Number of distinct product IDs"),
            (
                "seller_count",
                "INT64",
                "Number of distinct sellers fulfilling this order",
            ),
            ("product_revenue", "FLOAT64", "Sum of item prices (excl. freight)"),
            ("freight_revenue", "FLOAT64", "Sum of freight values"),
            ("total_order_value", "FLOAT64", "product_revenue + freight_revenue"),
            (
                "total_payment_value",
                "FLOAT64",
                "Sum of all payment records for this order",
            ),
            (
                "payment_installments",
                "INT64",
                "Maximum installments across payment methods",
            ),
            ("primary_payment_type", "STRING", "Payment type with the highest value"),
            ("review_score", "INT64", "Customer review score (1–5); NULL if no review"),
            (
                "actual_delivery_days",
                "INT64",
                "Days from purchase to delivery; NULL if not delivered",
            ),
            (
                "estimated_delivery_days",
                "INT64",
                "Days from purchase to estimated delivery",
            ),
            (
                "is_on_time",
                "BOOL",
                "TRUE if delivered_date ≤ estimated_delivery_date; NULL if undelivered",
            ),
            (
                "customer_state",
                "STRING",
                "Denormalised from DimCustomer for partition efficiency",
            ),
            ("customer_region", "STRING", "Denormalised region"),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    add_heading(doc, "5.2  FactOrderItems", level=2)
    doc.add_paragraph(
        "Grain: one row per (order_id, order_item_id) (112,650 rows)  |  Partition: shipping_limit_date  |  Cluster: seller_id, product_id"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("order_id", "STRING", "FK → FactOrders"),
            (
                "order_item_id",
                "INT64",
                "Item sequence number within the order (1, 2, 3, …)",
            ),
            ("product_id", "STRING", "FK → DimProduct"),
            ("seller_id", "STRING", "FK → DimSeller"),
            ("shipping_limit_date", "DATE", "Seller's shipping deadline"),
            ("item_price", "FLOAT64", "Selling price of the item"),
            ("freight_value", "FLOAT64", "Freight cost allocated to this item"),
            ("item_total", "FLOAT64", "item_price + freight_value"),
            ("product_category", "STRING", "Denormalised English category name"),
            ("product_weight_g", "INT64", "Product weight in grams"),
            ("seller_state", "STRING", "Denormalised seller state"),
            ("seller_region", "STRING", "Denormalised seller region"),
            ("purchase_date", "DATE", "Denormalised from FactOrders"),
            ("order_status", "STRING", "Denormalised from FactOrders"),
            ("customer_state", "STRING", "Denormalised from FactOrders"),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    add_heading(doc, "5.3  FactMarketingFunnel", level=2)
    doc.add_paragraph(
        "Grain: one row per mql_id (8,000 rows)  |  Partition: first_contact_date  |  Cluster: channel_code, business_segment"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("mql_id", "STRING", "Marketing Qualified Lead identifier (natural key)"),
            ("landing_page_id", "STRING", "Landing page where the lead originated"),
            ("channel_code", "STRING", "FK → DimMarketingChannel.channel_code"),
            ("first_contact_date", "DATE", "Date the lead first contacted Olist"),
            ("won_date", "DATE", "Date the deal was closed; NULL if not converted"),
            ("seller_id", "STRING", "FK → DimSeller (for converted leads only)"),
            ("business_segment", "STRING", "Business vertical of the converted seller"),
            ("lead_type", "STRING", "Lead type classification"),
            ("lead_behaviour_profile", "STRING", "Behavioural profile score"),
            ("business_type", "STRING", "Reseller, manufacturer, other"),
            (
                "declared_monthly_revenue",
                "FLOAT64",
                "Self-reported monthly revenue of the seller",
            ),
            ("is_converted", "BOOL", "TRUE if the lead became a seller"),
            (
                "days_to_close",
                "INT64",
                "Days from first_contact_date to won_date; NULL if not converted",
            ),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    add_heading(doc, "5.4  FactPayments", level=2)
    doc.add_paragraph(
        "Grain: one row per (order_id, payment_sequential) (103,886 rows)  |  Cluster: payment_type"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("order_id", "STRING", "FK → FactOrders"),
            (
                "payment_sequential",
                "INT64",
                "Payment sequence number (1 = primary; 2+ = secondary)",
            ),
            ("payment_type", "STRING", "FK → DimPaymentType.payment_type_code"),
            ("payment_installments", "INT64", "Number of instalments for this payment"),
            ("payment_value", "FLOAT64", "Payment amount"),
            ("purchase_date", "DATE", "Denormalised from FactOrders"),
            ("customer_state", "STRING", "Denormalised from FactOrders"),
            ("order_status", "STRING", "Denormalised from FactOrders"),
            ("dw_inserted_at", "TIMESTAMP", "Pipeline run timestamp"),
        ],
    )

    # ── Section 6: Mart Tables ─────────────────────────────────────────────────
    add_heading(doc, "6.  Mart Tables (olist_analytics_marts)")

    add_heading(doc, "6.1  mart_monthly_revenue", level=2)
    doc.add_paragraph(
        "Grain: (revenue_month, category) — ~300 rows  |  Audience: CEO, CFO, Finance team  |  Refresh: daily full"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("revenue_month", "DATE", "First day of the calendar month"),
            ("year / month", "INT64", "Year and month numbers"),
            ("category", "STRING", "Product category (English name or 'Unknown')"),
            ("product_revenue", "FLOAT64", "Item price revenue (excl. freight)"),
            ("freight_revenue", "FLOAT64", "Freight revenue"),
            ("total_revenue", "FLOAT64", "product_revenue + freight_revenue"),
            ("order_count", "INT64", "Distinct orders"),
            ("unique_customers", "INT64", "Distinct customer_unique_ids"),
            ("avg_item_price", "FLOAT64", "Average item price"),
            ("avg_order_value", "FLOAT64", "total_revenue / order_count"),
            (
                "prev_month_revenue",
                "FLOAT64",
                "Previous month's total_revenue (same category)",
            ),
            ("mom_growth_pct", "FLOAT64", "Month-over-month growth percentage"),
            (
                "ytd_revenue",
                "FLOAT64",
                "Year-to-date cumulative revenue (same category)",
            ),
        ],
    )

    add_heading(doc, "6.2  mart_customer_lifetime_value", level=2)
    doc.add_paragraph(
        "Grain: customer_unique_id — one row per unique customer with at least one delivered order  |  Audience: CMO, CRM, growth analytics"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("customer_unique_id", "STRING", "Stable customer identifier"),
            ("customer_state / customer_city", "STRING", "Customer location"),
            ("first_order_date", "DATE", "Date of first delivered order"),
            ("last_order_date", "DATE", "Date of most recent delivered order"),
            ("order_count", "INT64", "Total delivered orders"),
            (
                "total_product_spend",
                "FLOAT64",
                "Sum of item prices across all delivered orders",
            ),
            ("total_freight_spend", "FLOAT64", "Sum of freight values"),
            ("total_spend", "FLOAT64", "Total amount paid (product + freight)"),
            ("avg_order_value", "FLOAT64", "Average per-order spend"),
            ("customer_tenure_days", "INT64", "Days between first and last order"),
            ("recency_days", "INT64", "Days since last order (relative to 2018-12-31)"),
            (
                "r_score / f_score / m_score",
                "INT64",
                "RFM quartile scores (1=worst, 4=best)",
            ),
            (
                "rfm_segment",
                "STRING",
                "Champions, Loyal Customers, New Customers, Potential Loyalists, At Risk, Cant Lose Them, Sleeping Giants, Hibernating",
            ),
            ("value_tier", "STRING", "High Value, Medium Value, Low Value"),
            (
                "estimated_clv",
                "FLOAT64",
                "Proxy CLV = total_spend × (1 + 0.3 × min(order_count-1, 5))",
            ),
        ],
    )

    add_heading(doc, "6.3  mart_geo_performance", level=2)
    doc.add_paragraph(
        "Grain: (state_code, order_month) — incremental  |  Audience: COO, regional managers, expansion team"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("state_code", "STRING", "2-letter state code"),
            ("state_name / region", "STRING", "Full name and macro-region"),
            ("order_month", "DATE", "First day of the month"),
            ("total_orders", "INT64", "Orders from customers in this state this month"),
            ("unique_customers", "INT64", "Distinct customers"),
            ("total_revenue", "FLOAT64", "Sum of item + freight revenue"),
            ("avg_order_value", "FLOAT64", "Average per-order revenue"),
            ("revenue_per_customer", "FLOAT64", "total_revenue / unique_customers"),
            (
                "state_population",
                "INT64",
                "2018 Brazilian state population (hardcoded reference)",
            ),
            ("orders_per_1k_pop", "FLOAT64", "Market penetration proxy"),
            ("revenue_per_capita", "FLOAT64", "Revenue divided by state population"),
            ("is_frontier_market", "BOOL", "High e-commerce infrastructure state"),
        ],
    )

    add_heading(doc, "6.4  mart_seller_performance", level=2)
    doc.add_paragraph(
        "Grain: seller_id — one row per seller (all-time)  |  Audience: COO, marketplace operations, account management"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("seller_id", "STRING", "Unique seller identifier"),
            ("seller_state / seller_city / seller_region", "STRING", "Seller location"),
            (
                "seller_tier",
                "STRING",
                "Platinum (top 10%), Gold (11–30%), Silver (31–60%), Bronze (61–100%)",
            ),
            (
                "performance_score",
                "FLOAT64",
                "0–100 composite: 40% review score, 35% on-time delivery, 25% cancellation rate",
            ),
            ("total_orders", "INT64", "Distinct orders fulfilled"),
            ("total_items_sold", "INT64", "Total line items"),
            ("unique_products", "INT64", "Distinct products listed"),
            ("gross_merchandise_value", "FLOAT64", "Sum of item prices"),
            ("total_freight_charged", "FLOAT64", "Sum of freight revenue"),
            ("total_revenue", "FLOAT64", "GMV + freight"),
            ("avg_item_price", "FLOAT64", "Average item selling price"),
            (
                "on_time_delivery_pct",
                "FLOAT64",
                "% of deliveries within estimated date",
            ),
            ("avg_review_score", "FLOAT64", "Average customer review score"),
            ("cancellation_rate_pct", "FLOAT64", "% of orders cancelled"),
            ("first_sale_date / last_sale_date", "DATE", "Active selling window"),
            ("active_selling_days", "INT64", "Days between first and last sale"),
        ],
    )

    add_heading(doc, "6.5  mart_marketing_funnel", level=2)
    doc.add_paragraph(
        "Grain: (origin, cohort_month) — one row per channel per month  |  Audience: CMO, marketing operations, growth team"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            (
                "origin",
                "STRING",
                "Acquisition channel (paid_search, organic_search, social, email, etc.)",
            ),
            (
                "conversion_tier",
                "STRING",
                "High Converting (≥10%), Mid Converting (5–10%), Low Converting (<5%)",
            ),
            ("cohort_month", "DATE", "Month of first MQL contact"),
            ("total_leads", "INT64", "MQLs in this cohort"),
            ("converted_leads", "INT64", "Leads that became sellers"),
            ("dropped_leads", "INT64", "Leads that did not convert"),
            ("conversion_rate_pct", "FLOAT64", "converted_leads / total_leads × 100"),
            (
                "avg_days_to_close",
                "FLOAT64",
                "Average days from first contact to deal won",
            ),
            (
                "avg_declared_monthly_revenue",
                "FLOAT64",
                "Average self-reported monthly revenue of converted sellers",
            ),
            (
                "mom_lead_growth_pct",
                "FLOAT64",
                "Month-over-month change in lead volume",
            ),
            (
                "all_time_conversion_pct",
                "FLOAT64",
                "Channel's all-time conversion rate",
            ),
        ],
    )

    add_heading(doc, "6.6  mart_logistics_performance", level=2)
    doc.add_paragraph(
        "Grain: (customer_state, order_month) — incremental  |  Audience: COO, logistics/supply chain, carrier management"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("customer_state", "STRING", "2-letter state code"),
            ("region", "STRING", "Brazilian macro-region"),
            ("order_month", "DATE", "Month of purchase"),
            ("total_delivered_orders", "INT64", "Delivered orders in this state/month"),
            ("avg_delivery_days", "FLOAT64", "Mean delivery time"),
            ("median_delivery_days", "FLOAT64", "P50 delivery time"),
            ("p90_delivery_days", "FLOAT64", "90th percentile delivery time"),
            ("on_time_pct", "FLOAT64", "% delivered ≤ estimated date"),
            (
                "avg_delay_days_when_late",
                "FLOAT64",
                "Mean delay for orders that were late",
            ),
            (
                "avg_review_score",
                "FLOAT64",
                "Average review score for this state/month",
            ),
            (
                "avg_review_on_time / avg_review_late",
                "FLOAT64",
                "Average review for on-time vs late deliveries",
            ),
            (
                "review_score_delta_on_time_vs_late",
                "FLOAT64",
                "avg_review_on_time − avg_review_late",
            ),
            (
                "sla_performance_band",
                "STRING",
                "Excellent (≥95%), Good (≥90%), Needs Improvement (≥80%), Critical (<80%)",
            ),
            (
                "orders_0_5_days … orders_over_30_days",
                "INT64",
                "Delivery speed bucket counts",
            ),
        ],
    )

    add_heading(doc, "6.7  mart_product_performance", level=2)
    doc.add_paragraph(
        "Grain: category_english — one row per product category (all-time)  |  Audience: category managers, merchandising, CPO"
    )
    add_column_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("category_english", "STRING", "English product category name"),
            (
                "performance_tier",
                "STRING",
                "Tier 1 — Core, Tier 2 — Growth, Tier 3 — Niche, Tier 4 — Tail",
            ),
            ("gmv_quartile", "INT64", "GMV rank quartile (1=top 25%)"),
            ("total_skus", "INT64", "Distinct products in this category"),
            ("total_orders", "INT64", "Orders containing this category"),
            ("total_units_sold", "INT64", "Total units sold"),
            ("unique_customers", "INT64", "Distinct buyers"),
            ("category_gross_revenue", "FLOAT64", "Sum of item prices"),
            ("category_freight_revenue", "FLOAT64", "Sum of freight"),
            ("total_category_revenue", "FLOAT64", "Gross + freight"),
            (
                "avg_item_price / median_item_price",
                "FLOAT64",
                "Mean and P50 item price",
            ),
            ("revenue_share_pct", "FLOAT64", "Category's share of total platform GMV"),
            (
                "cumulative_revenue_pct",
                "FLOAT64",
                "Cumulative revenue % (Lorenz curve value)",
            ),
            ("avg_review_score", "FLOAT64", "Average customer satisfaction score"),
            ("positive_review_rate_pct", "FLOAT64", "% reviews scoring 4 or 5"),
            ("on_time_delivery_pct", "FLOAT64", "% of orders delivered on time"),
            ("avg_units_per_order", "FLOAT64", "Average basket size in units"),
            ("revenue_per_order", "FLOAT64", "Average revenue per order"),
            ("first_sale_date / last_sale_date", "DATE", "Active selling window"),
        ],
    )

    # ── Section 7: Business Glossary ───────────────────────────────────────────
    add_heading(doc, "7.  Business Glossary")
    add_column_table(
        doc,
        ["Term", "Definition"],
        [
            (
                "GMV",
                "Gross Merchandise Value — total value of goods sold. Computed as SUM(item_price) across all non-cancelled orders.",
            ),
            ("AOV", "Average Order Value — total_order_value / order_count"),
            (
                "CLV",
                "Customer Lifetime Value — estimated total revenue from a customer. Phase 1: proxy formula. Phase 2: BG/NBD model.",
            ),
            (
                "RFM",
                "Recency × Frequency × Monetary — customer segmentation framework using three purchase behaviour dimensions.",
            ),
            (
                "MQL",
                "Marketing Qualified Lead — a prospective seller who has expressed intent to join the Olist platform.",
            ),
            (
                "On-time delivery",
                "delivered_date ≤ estimated_delivery_date as shown at the time of purchase.",
            ),
            (
                "customer_id",
                "Order-level customer identifier. One person placing two orders gets two customer_ids.",
            ),
            (
                "customer_unique_id",
                "Person-level stable identifier. Multiple orders by the same person share one unique_id.",
            ),
            (
                "is_frontier_market",
                "FALSE for 10 low-penetration states: AC, AL, AM, AP, MA, PB, RN, RO, RR, TO. TRUE for the other 17.",
            ),
            (
                "Seller tier",
                "Platinum (top 10% by GMV), Gold (11–30%), Silver (31–60%), Bronze (61–100%).",
            ),
            (
                "Performance score",
                "Composite 0–100 metric per seller: 40% review + 35% on-time delivery + 25% cancellation rate.",
            ),
        ],
    )

    doc.save("docs/OLP-DD-001_Data_Dictionary.docx")
    print("Saved docs/OLP-DD-001_Data_Dictionary.docx")


if __name__ == "__main__":
    build()
