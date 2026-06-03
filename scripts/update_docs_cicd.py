"""
scripts/update_docs_cicd.py
============================
Appends CI/CD framework content to:
  - docs/OLP-OPS-001_Operations_Manual.docx
  - docs/OLP-TDD-001_Technical_Design_Document.docx
  - docs/OLP-SRS-001_Software_Requirements_Specification.docx

Run:
  conda activate olist
  cd olist_platform
  python scripts/update_docs_cicd.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


# ── Helpers ────────────────────────────────────────────────────────────────────


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    sizes = {1: 16, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for run in hdr.cells[i].paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()


def add_code(doc, text):
    p = doc.add_paragraph(style="No Spacing")
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def section_exists(doc, keyword):
    return any(keyword.lower() in p.text.lower() for p in doc.paragraphs)


# ── OLP-OPS-001: Operations Manual ────────────────────────────────────────────


def update_ops_manual():
    path = DOCS / "OLP-OPS-001_Operations_Manual.docx"
    doc = Document(path)

    if section_exists(doc, "GitHub Actions CI/CD"):
        print(f"OPS Manual: CI/CD section already exists — skipping.")
        return

    doc.add_page_break()
    add_heading(doc, "GitHub Actions CI/CD — Operations Runbook", level=1)

    add_para(
        doc,
        (
            "This section covers day-to-day operations of the GitHub Actions CI/CD framework. "
            "Four workflows are live in .github/workflows/. This runbook covers how to monitor, "
            "re-run, debug, and extend them."
        ),
    )

    add_heading(doc, "Workflow Overview", level=2)
    add_table(
        doc,
        ["Workflow", "Trigger", "Purpose"],
        [
            [
                "ci.yml",
                "Push / PR to main",
                "Lint · dbt parse · security scan. No BigQuery needed.",
            ],
            [
                "dbt-ci.yml",
                "PR touching dbt_project/**",
                "dbt compile + slim build (state:modified+)",
            ],
            [
                "pipeline-daily.yml",
                "Cron 06:00 UTC / workflow_dispatch",
                "Full ELT: DQ → dbt staging → marts → integration tests → Slack",
            ],
            [
                "deploy-dashboard.yml",
                "Push to dashboard/** / workflow_dispatch",
                "Deploy dashboard/index.html to GitHub Pages",
            ],
        ],
    )

    add_heading(doc, "How to Re-run a Failed Workflow", level=2)
    add_para(doc, "1.  Go to the repository on GitHub.")
    add_para(doc, "2.  Click the Actions tab.")
    add_para(doc, "3.  Find the failed run in the list.")
    add_para(doc, "4.  Click into the run → click Re-run failed jobs (top right).")
    add_para(
        doc,
        "5.  If the job uses environment: production or bigquery-ci and requires approval, click Review deployments → Approve and deploy.",
    )

    add_heading(doc, "How to Trigger the Daily Pipeline Manually", level=2)
    add_para(
        doc,
        "Actions tab → Daily Pipeline — Full Olist ELT → Run workflow dropdown → set full_pipeline:",
    )
    add_para(
        doc,
        "  false (default) — runs dbt staging + marts + integration tests only (data already in BigQuery).",
    )
    add_para(
        doc,
        "  true — downloads CSVs from GCS, runs Python DQ validation, Meltano ELT, then dbt.",
    )

    add_heading(doc, "Troubleshooting Common Failures", level=2)
    add_table(
        doc,
        ["Error", "Root Cause", "Fix"],
        [
            [
                "dbt found 1 package(s) but 0 installed",
                "dbt deps not run before dbt command",
                "Add: dbt deps --profiles-dir profiles step before the failing dbt step",
            ],
            [
                "401 Anonymous caller — storage.objects.list denied",
                "gsutil not authenticated (setup-gcloud@v2 missing)",
                "Add: uses: google-github-actions/setup-gcloud@v2 after auth@v2",
            ],
            [
                "403 <service-account> does not have storage.objects.list",
                "Service account missing objectAdmin on GCS bucket",
                "Run: gsutil iam ch serviceAccount:<email>:objectAdmin gs://olist-bucket-01",
            ],
            [
                "InvalidUrlError: Found gs:///dbt-state/manifest.json",
                "GCS_BUCKET_NAME variable not set (empty = three slashes)",
                "Settings → Secrets → Variables tab → add GCS_BUCKET_NAME = olist-bucket-01",
            ],
            [
                "Downstream jobs skipped despite upstream success",
                "always() missing — GitHub propagates skipped transitively",
                "Add always() && prefix to every if: condition in the job chain",
            ],
        ],
    )

    add_heading(doc, "GitHub Secrets and Variables", level=2)
    add_para(doc, "Navigate to: Settings → Secrets and variables → Actions")
    add_table(
        doc,
        ["Name", "Type", "Value"],
        [
            [
                "GCP_SERVICE_ACCOUNT_KEY",
                "Secret",
                "Full JSON of olist-analytics-gcp-key.json",
            ],
            ["SLACK_WEBHOOK_URL", "Secret", "Slack incoming webhook URL (optional)"],
            ["GCP_PROJECT_ID", "Variable (not Secret)", "olist-analytics-01"],
            ["GCS_BUCKET_NAME", "Variable (not Secret)", "olist-bucket-01"],
        ],
    )

    add_heading(doc, "GitHub Environments Required", level=2)
    add_table(
        doc,
        ["Environment", "Used by", "Notes"],
        [
            [
                "bigquery-ci",
                "dbt-ci.yml",
                "Gates GCP_SERVICE_ACCOUNT_KEY access on PRs",
            ],
            [
                "production",
                "pipeline-daily.yml",
                "Gates production BigQuery access on daily runs",
            ],
            [
                "github-pages",
                "deploy-dashboard.yml",
                "Auto-created on first successful Pages deploy",
            ],
        ],
    )

    doc.save(path)
    print(f"OPS Manual: CI/CD section appended → {path}")


# ── OLP-TDD-001: Technical Design Document ────────────────────────────────────


def update_tdd():
    path = DOCS / "OLP-TDD-001_Technical_Design_Document.docx"
    doc = Document(path)

    if section_exists(doc, "GitHub Actions CI/CD"):
        print(f"TDD: CI/CD section already exists — skipping.")
        return

    doc.add_page_break()
    add_heading(doc, "GitHub Actions CI/CD — Technical Design", level=1)

    add_para(
        doc,
        (
            "The CI/CD framework uses four GitHub Actions workflows. Each workflow has a distinct "
            "responsibility and trigger. They complement Dagster (which handles runtime data pipeline "
            "orchestration) by covering the software engineering lifecycle: code quality, SQL validation, "
            "and dashboard deployment."
        ),
    )

    add_heading(doc, "Architecture: GitHub Actions vs Dagster", level=2)
    add_table(
        doc,
        ["Concern", "GitHub Actions", "Dagster"],
        [
            ["Python linting / formatting", "ci.yml — ruff on every push", "—"],
            ["dbt SQL validation", "dbt-ci.yml — dbt compile on PR", "—"],
            ["Security scanning", "ci.yml — bandit on every push", "—"],
            [
                "Scheduled data pipeline",
                "pipeline-daily.yml — backup/fallback",
                "Primary — 06:00 UTC, asset graph, retries",
            ],
            [
                "Pipeline observability",
                "Flat pass/fail logs only",
                "Rich UI — run history, retries, asset lineage",
            ],
            ["Dashboard deployment", "deploy-dashboard.yml — GitHub Pages", "—"],
            [
                "Manual pipeline trigger",
                "workflow_dispatch",
                "Materialise all in Dagster UI",
            ],
        ],
    )

    add_heading(doc, "Slim CI Design (dbt-ci.yml)", level=2)
    add_para(
        doc,
        (
            "dbt-ci.yml implements the slim CI pattern to avoid running all 16 dbt models on every PR. "
            "After each successful compile, the manifest.json is stored in GCS. The next PR downloads "
            "this manifest and runs dbt build --select state:modified+ to build and test only the "
            "changed models and their downstream dependents."
        ),
    )
    add_para(
        doc,
        "First PR ever: no prior manifest → falls back to full staging build → uploads first manifest to GCS.",
    )
    add_para(
        doc,
        "Subsequent PRs: downloads prior manifest → runs state:modified+ (changed models only) → uploads updated manifest.",
    )

    add_heading(doc, "Concurrency Strategy", level=2)
    add_table(
        doc,
        ["Workflow", "cancel-in-progress", "Reason"],
        [
            [
                "ci.yml",
                "true (feature branches), false (main)",
                "New push supersedes old check on same branch; main validates every commit",
            ],
            [
                "dbt-ci.yml",
                "true",
                "New PR commit invalidates previous CI run entirely",
            ],
            [
                "pipeline-daily.yml",
                "false",
                "Never cancel mid-run — partial BigQuery state is worse than a queued run",
            ],
            [
                "deploy-dashboard.yml",
                "true",
                "Newer deploy supersedes older — no partial state risk with static files",
            ],
        ],
    )

    add_heading(doc, "Key Implementation Patterns", level=2)
    add_para(
        doc,
        "1. dbt deps before every dbt command — dbt_packages/ is gitignored, packages must be installed on every runner.",
    )
    add_para(
        doc,
        "2. setup-gcloud@v2 after auth@v2 — required for gsutil to use authenticated credentials.",
    )
    add_para(
        doc,
        "3. always() on downstream jobs — prevents transitive skip propagation through conditional job chains.",
    )
    add_para(
        doc,
        "4. GCS bucket IAM — service account needs objectAdmin role granted once: gsutil iam ch ...",
    )
    add_para(
        doc,
        "5. Variables vs Secrets — non-sensitive config (GCP_PROJECT_ID, GCS_BUCKET_NAME) goes in Variables tab, not Secrets.",
    )

    doc.save(path)
    print(f"TDD: CI/CD section appended → {path}")


# ── OLP-SRS-001: Software Requirements Specification ──────────────────────────


def update_srs():
    path = DOCS / "OLP-SRS-001_Software_Requirements_Specification.docx"
    doc = Document(path)

    if section_exists(doc, "GitHub Actions CI/CD"):
        print(f"SRS: CI/CD section already exists — skipping.")
        return

    doc.add_page_break()
    add_heading(doc, "GitHub Actions CI/CD — Requirements", level=1)

    add_heading(doc, "Functional Requirements", level=2)
    add_table(
        doc,
        ["ID", "Requirement", "Implemented by"],
        [
            [
                "CI-01",
                "Every push to any branch must trigger automated Python linting and formatting checks",
                "ci.yml — ruff",
            ],
            [
                "CI-02",
                "Every push must trigger a dbt SQL syntax validation without requiring BigQuery credentials",
                "ci.yml — dbt parse",
            ],
            [
                "CI-03",
                "Every push must trigger a Python security scan",
                "ci.yml — bandit",
            ],
            [
                "CI-04",
                "Every PR touching dbt_project/** must validate SQL against live BigQuery",
                "dbt-ci.yml — dbt compile",
            ],
            [
                "CI-05",
                "PRs must only rebuild and test the changed dbt models and their downstream dependents",
                "dbt-ci.yml — state:modified+",
            ],
            [
                "CI-06",
                "The full ELT pipeline must run automatically at 06:00 UTC daily as a fallback",
                "pipeline-daily.yml — cron",
            ],
            [
                "CI-07",
                "The pipeline must be manually triggerable with a full re-ingestion option",
                "pipeline-daily.yml — workflow_dispatch",
            ],
            [
                "CI-08",
                "Dashboard updates must be automatically deployed to GitHub Pages on merge to main",
                "deploy-dashboard.yml",
            ],
            [
                "CI-09",
                "Pipeline failures must send a Slack notification",
                "pipeline-daily.yml — Slack webhook",
            ],
        ],
    )

    add_heading(doc, "Non-Functional Requirements", level=2)
    add_table(
        doc,
        ["ID", "Requirement", "How met"],
        [
            [
                "NFR-CI-01",
                "No credentials may appear in YAML workflow files",
                "google-github-actions/auth@v2 + GitHub Secrets",
            ],
            [
                "NFR-CI-02",
                "ci.yml must require no BigQuery credentials",
                "dbt parse uses dummy credentials; no real connection",
            ],
            [
                "NFR-CI-03",
                "Running pipeline must never be cancelled mid-run",
                "concurrency: cancel-in-progress: false on pipeline-daily.yml",
            ],
            [
                "NFR-CI-04",
                "CI/CD must not replace Dagster scheduling",
                "pipeline-daily.yml is a fallback; Dagster remains primary",
            ],
        ],
    )

    doc.save(path)
    print(f"SRS: CI/CD section appended → {path}")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    update_ops_manual()
    update_tdd()
    update_srs()
    print("\nAll docx files updated.")
