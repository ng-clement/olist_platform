"""Rebuild presentation_transcript.docx from presentation_transcript.md."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x0D, 0x3B, 0x66)
TEAL = RGBColor(0x1C, 0x72, 0x93)
GRAY = RGBColor(0x55, 0x55, 0x55)

MD = Path("docs/presentation_transcript.md").read_text()


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    # Title block
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("OLIST ANALYTICS PLATFORM")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Presentation Transcript & Speaker Notes")
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = TEAL

    meta = [
        "Presentation: Olist Analytics Platform",
        "Duration: 10 minutes presentation + 5 minutes Q&A",
        "Audience: Business executives (CEO, CFO, CMO, COO) and technical leadership (CTO, VP Engineering)",
        "Structure: NTU Coaching 2.3 — Executive Summary → Introduction & Context → Methodology & Data → Results & Insights → Strategic Recommendations → Conclusion",
    ]
    for m in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY

    doc.add_paragraph()

    # Parse the markdown into slides
    lines = MD.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Slide heading
        if line.startswith("## SLIDE "):
            h = doc.add_heading(line[3:], level=1)
            for run in h.runs:
                run.font.color.rgb = NAVY
                run.font.size = Pt(14)
            i += 1
            continue

        # Sub-section bold heading
        if line.startswith("**") and line.endswith("**") and line.count("**") == 2:
            p = doc.add_paragraph()
            r = p.add_run(line.strip("* "))
            r.bold = True
            r.font.color.rgb = TEAL
            r.font.size = Pt(10.5)
            i += 1
            continue

        # Q&A Section heading
        if line.startswith("## Q&A"):
            doc.add_page_break()
            h = doc.add_heading("Q&A Preparation", level=1)
            for run in h.runs:
                run.font.color.rgb = NAVY
            i += 1
            continue

        # Bold Q&A question
        if line.startswith("**Q:"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip("* "))
            r.bold = True
            r.font.color.rgb = TEAL
            r.font.size = Pt(10)
            i += 1
            continue

        # Answer line
        if line.startswith("A:"):
            p = doc.add_paragraph(line[2:].strip())
            p.runs[0].font.size = Pt(9.5) if p.runs else None
            i += 1
            continue

        # HR separator
        if line.strip() == "---":
            p = doc.add_paragraph("─" * 80)
            p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            p.runs[0].font.size = Pt(8)
            i += 1
            continue

        # Quoted speech paragraph
        if line.startswith('"') or (line and lines[i - 1:i] and '"' in line):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.italic = True
            i += 1
            continue

        # Bold inline text in paragraph (partial bold)
        if "**" in line and line.strip():
            p = doc.add_paragraph()
            parts = line.split("**")
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.bold = (j % 2 == 1)
                run.font.size = Pt(10)
            i += 1
            continue

        # Normal non-empty paragraph
        if line.strip():
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # Empty line → paragraph break already handled by empty add
        i += 1

    doc.save("docs/presentation_transcript.docx")
    print("Saved docs/presentation_transcript.docx")


if __name__ == "__main__":
    build()
